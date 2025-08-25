#!/usr/bin/python
# -*- coding: latin-1 -*-

from docx import Document
import os

def DC1_file(dc):
    replacements = {'ACHTEUR':dc['DC1.1'],
                    "CONSULTATION_OBJECT":dc['DC1.2'],
                    "XY":dc['DC1.3']
                    } 
    
    # Load the document
    doc = Document("/home/aymen/SEF/examples_files/DC1-2019.docx")

    # Function to replace text in paragraphs
    def replace_text_in_paragraphs(paragraphs, replacements):
        for paragraph in paragraphs:
            for word, replacement in replacements.items():
              for run in paragraph.runs:  # Itère sur chaque "run" du paragraphe
                if word in run.text:  # Si le mot à remplacer est dans ce "run"
                    run.text = run.text.replace(word, replacement)  # Remplace le texte
                    run.bold = True  # Met le texte en gras
                    
                #if word in paragraph.text:
                    #paragraph.text = paragraph.text.replace(word, replacement)
                    #paragraph.bold = True

    # Replace words in body paragraphs
    replace_text_in_paragraphs(doc.paragraphs, replacements)

    # Replace words in tables
    #for table in doc.tables:
    #    for row in table.rows:
    #        for cell in row.cells:
    #            replace_text_in_paragraphs(cell.paragraphs, replacements)

    # Replace words in footers
    #for section in doc.sections:
    #    for i in section.footer.tables:
    #      for row in i.rows:
    #        for cell in row.cells:
    #           replace_text_in_paragraphs(cell.paragraphs, replacements)
    
    # Save the modified document
    doc.save('/home/aymen/SEF/output_file/DC1_2019.docx')
    #os.system(f'abiword --to pdf "/home/aymen/SEF/output_file/DC1_2019.docx"')
    print("DC1 created successfully.")

def DC2_file(dc):
    replacements = {'ACHTEUR_':dc['DC1.1'],
                    "OBJET_PROJET":dc['DC1.2']
                    } 
    
    # Load the document
    doc = Document("/home/aymen/SEF/examples_files/DC2-2019.docx")

    # Function to replace text in paragraphs
    def replace_text_in_paragraphs(paragraphs, replacements):
        for paragraph in paragraphs:
            for word, replacement in replacements.items():
              for run in paragraph.runs:  # Itère sur chaque "run" du paragraphe
                if word in run.text:  # Si le mot à remplacer est dans ce "run"
                    run.text = run.text.replace(word, replacement)  # Remplace le texte
                    run.bold = True  # Met le texte en gras
                    
                #if word in paragraph.text:
                    #paragraph.text = paragraph.text.replace(word, replacement)
                    #paragraph.bold = True

    # Replace words in body paragraphs
    replace_text_in_paragraphs(doc.paragraphs, replacements)

    # Replace words in tables
    #for table in doc.tables:
    #    for row in table.rows:
    #        for cell in row.cells:
    #            replace_text_in_paragraphs(cell.paragraphs, replacements)

    # Replace words in footers
    #for section in doc.sections:
    #    for i in section.footer.tables:
    #      for row in i.rows:
    #        for cell in row.cells:
    #           replace_text_in_paragraphs(cell.paragraphs, replacements)
    
    # Save the modified document
    doc.save('/home/aymen/SEF/output_file/DC2_2019.docx')
    #os.system(f'abiword --to pdf "/home/aymen/SEF/output_file/DC1_2019.docx"')
    print("DC1 created successfully.")



if __name__ == "__main__":
    # Define your file path and the words to replace
    
    DC1_file()