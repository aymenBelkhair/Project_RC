#!/usr/bin/python
# -*- coding: latin-1 -*-

from docx import Document

def DC1_file(replacements):
    # Load the document
    doc = Document("/home/aymen/SEF/examples_files/DC1-2019.docx")

    # Function to replace text in paragraphs
    def replace_text_in_paragraphs(paragraphs, replacements):
        for paragraph in paragraphs:
            for word, replacement in replacements.items():
                if word in paragraph.text:
                    paragraph.text = paragraph.text.replace(word, replacement)

    # Replace words in body paragraphs
    replace_text_in_paragraphs(doc.paragraphs, replacements)

    # Replace words in tables
    #for table in doc.tables:
    #    for row in table.rows:
    #        for cell in row.cells:
    #            replace_text_in_paragraphs(cell.paragraphs, replacements)

    # Replace words in footers
    for section in doc.sections:
        for i in section.footer.tables:
          for row in i.rows:
            for cell in row.cells:
               replace_text_in_paragraphs(cell.paragraphs, replacements)
    
    # Save the modified document
    doc.save('DC1_2019.docx')
    print('done')



if __name__ == "__main__":
    # Define your file path and the words to replace
    replacements = {'ACHTEUR':'Answer q1',
                    "CONSULTATION_OBJECT":"Answer q2",
                    "lot n°":"lot n° : x ",
                    "Projet":"++++"
                    }
    DC1_file(replacements)